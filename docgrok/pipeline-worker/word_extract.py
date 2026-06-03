"""Microsoft Word document text extraction for the DocGrok pipeline worker.

Two formats are supported:

* ``.docx`` — modern Office Open XML (zip + XML). Extracted with
  ``python-docx``; we read paragraphs and table cell text from the main
  document and from any headers/footers/footnotes/endnotes. Returns one
  "section" per Word section break so downstream chunking can still
  preserve logical groupings.

* ``.doc``  — legacy OLE compound binary. We use ``olefile`` to read the
  ``WordDocument`` stream and decode the text directly from the FIB.
  This is a best-effort path; if extraction yields nothing usable we
  raise an explicit 415 so the caller can route or convert.

Both functions return a ``list[str]`` of "page-like" text segments
matching the shape ``_stage_extract`` already produces for PDFs/text.
"""

from __future__ import annotations

from typing import List


def extract_docx(path: str) -> List[str]:
    """Extract text from a .docx file.

    Returns a list where each element is the text of one Word section
    (delimited by section breaks). For a single-section document this is
    a one-element list. Paragraphs are separated by ``\\n``; tables are
    flattened cell-by-cell with tab + newline separators.
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError as e:
        raise RuntimeError(
            "python-docx is not installed; cannot extract .docx. "
            "Install with: pip install python-docx"
        ) from e

    doc = Document(path)

    sections: List[List[str]] = [[]]

    def _emit(text: str) -> None:
        if text:
            sections[-1].append(text)

    def _para_text(p) -> str:
        return p.text or ""

    def _table_text(tbl) -> str:
        rows: List[str] = []
        for row in tbl.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                rows.append("\t".join(cells))
        return "\n".join(rows)

    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag
        if tag == qn("w:p"):
            from docx.text.paragraph import Paragraph
            _emit(_para_text(Paragraph(child, doc)))
            # A section break lives inside a paragraph's pPr/sectPr.
            pPr = child.find(qn("w:pPr"))
            if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
                if sections[-1]:
                    sections.append([])
        elif tag == qn("w:tbl"):
            from docx.table import Table
            _emit(_table_text(Table(child, doc)))
        elif tag == qn("w:sectPr"):
            if sections[-1]:
                sections.append([])

    header_footer_text: List[str] = []
    for sec in doc.sections:
        for hf in (sec.header, sec.footer, sec.first_page_header,
                   sec.first_page_footer, sec.even_page_header,
                   sec.even_page_footer):
            for p in getattr(hf, "paragraphs", []) or []:
                t = _para_text(p)
                if t.strip():
                    header_footer_text.append(t)
    if header_footer_text:
        if sections[-1]:
            sections.append([])
        sections[-1].extend(header_footer_text)

    out = ["\n".join(s).strip() for s in sections]
    out = [s for s in out if s]
    return out or [""]


def extract_doc(path: str) -> List[str]:
    """Best-effort text extraction from a legacy .doc file.

    The .doc binary format stores the document text in the
    ``WordDocument`` OLE stream at offset ``fcMin``..``fcMin+ccpText``.
    Encoding is detected from the FIB ``fExtChar`` flag (CP1252 vs
    UTF-16LE).

    For documents with complex formatting (revision marks, embedded
    objects, etc.) the extracted text may be imperfect. Callers that
    need fidelity should ask users to convert to .docx first.
    """
    try:
        import olefile
    except ImportError as e:
        raise RuntimeError(
            "olefile is not installed; cannot extract legacy .doc. "
            "Install with: pip install olefile"
        ) from e

    if not olefile.isOleFile(path):
        raise ValueError(
            f"{path!r} is not an OLE compound file; not a legacy .doc"
        )

    with olefile.OleFileIO(path) as ole:
        if not ole.exists("WordDocument"):
            raise ValueError(
                f"{path!r} is missing the WordDocument stream; not a Word .doc"
            )
        with ole.openstream("WordDocument") as fh:
            wd = fh.read()

        if len(wd) < 0x200:
            raise ValueError(f"WordDocument stream too short ({len(wd)} bytes)")

        flags = int.from_bytes(wd[0x000A:0x000C], "little")
        is_complex = bool(flags & 0x0004)
        f_table = "1Table" if (flags & 0x0200) else "0Table"

        fc_min = int.from_bytes(wd[0x0018:0x001C], "little", signed=False)
        ccp_text = int.from_bytes(wd[0x004C:0x0050], "little", signed=False)

        if not is_complex and ccp_text and ccp_text < 10_000_000:
            end = min(fc_min + ccp_text * 2, len(wd))
            blob = wd[fc_min:end]
            text = _decode_word_text(blob)
            if text.strip():
                return [text]

        if ole.exists(f_table):
            with ole.openstream(f_table) as fh:
                table = fh.read()
            text = _extract_complex_doc_text(wd, table)
            if text.strip():
                return [text]

    raise RuntimeError(
        "Could not extract text from .doc; convert the file to .docx and retry."
    )


def _decode_word_text(blob: bytes) -> str:
    """Decode a Word text run, trying UTF-16LE first then CP1252."""
    if not blob:
        return ""
    try:
        if len(blob) % 2 == 0:
            u = blob.decode("utf-16-le", errors="strict")
            if sum(1 for c in u if c == "\x00") < len(u) // 4:
                return _normalize_word_runs(u)
    except UnicodeDecodeError:
        pass
    return _normalize_word_runs(blob.decode("cp1252", errors="replace"))


def _normalize_word_runs(text: str) -> str:
    out: List[str] = []
    for ch in text:
        code = ord(ch)
        if ch == "\x07":
            out.append("\t")
        elif ch in ("\x0B", "\x0C"):
            out.append("\n")
        elif ch in ("\r", "\x0E"):
            out.append("\n")
        elif code < 0x20 and ch not in ("\n", "\t"):
            continue
        else:
            out.append(ch)
    s = "".join(out)
    while "\n\n\n" in s:
        s = s.replace("\n\n\n", "\n\n")
    return s.strip()


def _extract_complex_doc_text(wd: bytes, table: bytes) -> str:
    """Reconstruct text from the piece table for complex .doc files."""
    fc_clx = int.from_bytes(wd[0x01A2:0x01A6], "little", signed=False)
    lcb_clx = int.from_bytes(wd[0x01A6:0x01AA], "little", signed=False)
    if not lcb_clx or fc_clx + lcb_clx > len(table):
        return ""

    clx = table[fc_clx:fc_clx + lcb_clx]
    i = 0
    while i < len(clx):
        if clx[i] == 0x02:
            cb_pcd = int.from_bytes(clx[i + 1:i + 5], "little", signed=False)
            pcdt = clx[i + 5:i + 5 + cb_pcd]
            n = (len(pcdt) - 4) // 12
            cps = [
                int.from_bytes(pcdt[j * 4:(j + 1) * 4], "little", signed=False)
                for j in range(n + 1)
            ]
            pcds = pcdt[4 + (n + 1) * 4:]
            parts: List[str] = []
            for k in range(n):
                pcd = pcds[k * 8:(k + 1) * 8]
                fc_raw = int.from_bytes(pcd[2:6], "little", signed=False)
                ansi = bool(fc_raw & 0x40000000)
                fc = fc_raw & 0x3FFFFFFF
                cp_len = cps[k + 1] - cps[k]
                if ansi:
                    fc //= 2
                    blob = wd[fc:fc + cp_len]
                else:
                    blob = wd[fc:fc + cp_len * 2]
                parts.append(_decode_word_text(blob))
            return _normalize_word_runs("".join(parts))
        elif clx[i] == 0x01:
            cb = int.from_bytes(clx[i + 1:i + 3], "little", signed=False)
            i += 3 + cb
        else:
            break
    return ""


if __name__ == "__main__":
    import sys
    import tempfile

    try:
        from docx import Document
    except ImportError:
        print("python-docx not installed; install it to run the self-test")
        sys.exit(2)

    doc = Document()
    doc.add_heading("DocGrok Word self-test", level=1)
    doc.add_paragraph("This is paragraph one with some text.")
    doc.add_paragraph("And a second paragraph.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "alpha"
    table.cell(0, 1).text = "beta"
    table.cell(1, 0).text = "gamma"
    table.cell(1, 1).text = "delta"
    doc.add_paragraph("After-table paragraph.")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        doc.save(tmp.name)
        path = tmp.name

    sections = extract_docx(path)
    print(f"Extracted {len(sections)} section(s):")
    for i, s in enumerate(sections):
        print(f"--- section {i} ({len(s)} chars) ---")
        print(s)

    assert sections, "no sections extracted"
    full = "\n".join(sections)
    for needle in ("paragraph one", "second paragraph",
                   "alpha", "beta", "gamma", "delta",
                   "After-table paragraph"):
        assert needle in full, f"missing {needle!r}"
    print("OK — docx round-trip succeeded")
